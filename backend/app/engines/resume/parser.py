"""
Document parsing: extract clean, layout-aware text from uploaded resume
PDF/DOCX files.

Primary engine: PyMuPDF (fitz) — the best open-source PDF text extractor.
It preserves reading order and spacing far better than pypdf's naive
extract_text(), which is what caused fields to merge with no separation
(e.g. "phone+91.../envel⌢peabhiram...@gmail.com/linkedinLinkedIn").

Strategy for PDFs:
  * Try fitz first (get_text("text") with block sort → natural reading order).
  * Post-process: normalise weird Unicode artefacts, collapse excessive blank
    lines, and ensure sections/words are separated by real spaces/newlines.
  * Fall back to pypdf only if fitz is unavailable.
"""

import io
import re
import unicodedata


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """Return cleaned text from a PDF or DOCX resume."""
    fname = (filename or "").lower()
    if fname.endswith(".pdf"):
        return _parse_pdf(file_bytes)
    elif fname.endswith(".docx"):
        return _parse_docx(file_bytes)
    else:
        raise ValueError("Unsupported file type. Upload a .pdf or .docx resume.")


def _parse_pdf(file_bytes: bytes) -> str:
    # Prefer PyMuPDF (layout-aware, best spacing/reading-order).
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        try:
            pages = []
            for page in doc:
                # Use word bounding boxes (not "text" mode) so tokens like
                # "Node.js", "gmail.com", "C++" stay intact instead of being
                # split as "Node . js" / "gmail . com".
                words = page.get_text("words")  # (x0,y0,x1,y1,word,block,line,wordno)
                lines: dict[float, list[tuple[float, str]]] = {}
                for w in words:
                    yb = round(w[1] / 4) * 4  # bucket by ~4px line height
                    lines.setdefault(yb, []).append((w[0], w[4]))
                page_lines = [
                    " ".join(tok for _, tok in sorted(lines[yb], key=lambda t: t[0]))
                    for yb in sorted(lines)
                ]
                pages.append("\n".join(page_lines))
            return _clean("\n".join(pages))
        finally:
            doc.close()
    except ImportError:
        pass
    # Fallback: pypdf (weaker spacing, but better than nothing).
    return _parse_pdf_pypdf(file_bytes)


def _parse_pdf_pypdf(file_bytes: bytes) -> str:
    try:
        import pypdf
    except ImportError:
        raise RuntimeError("No PDF parser available. Install PyMuPDF or pypdf.")
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    parts = [page.extract_text() or "" for page in reader.pages]
    return _clean("\n".join(parts))


def _parse_docx(file_bytes: bytes) -> str:
    try:
        import docx
    except ImportError:
        raise RuntimeError("python-docx is not installed. pip install python-docx")
    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _clean("\n".join(parts))


def _clean(text: str) -> str:
    """Normalise extraction artefacts so downstream parsing is reliable.

    NOTE: we deliberately do NOT insert spaces inside words (e.g. between a
    lowercase and an uppercase letter). The extractor already returns clean
    tokens via word bounding boxes; adding spaces there would re-break
    "Node.js" -> "Node . js". We only fix encoding/whitespace noise.
    """
    if not text:
        return ""
    # Unicode NFKC: fold full-width chars, strange variants, etc.
    text = unicodedata.normalize("NFKC", text)
    # Replace common non-breaking / zero-width artefacts.
    text = text.replace("\u00a0", " ").replace("\u200b", "").replace("\u200c", "")
    text = text.replace("&nbsp;", " ")
    text = text.replace("\u2022", "• ")
    # Drop control chars except newline/tab.
    text = "".join(ch if ch in "\n\t" or ord(ch) >= 32 else " " for ch in text)
    # Collapse 3+ spaces to one.
    text = re.sub(r"[ \t]{3,}", " ", text)
    # Collapse 3+ blank lines to a single blank line.
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Strip trailing whitespace per line, drop fully-empty edges.
    lines = [ln.rstrip() for ln in text.split("\n")]
    return "\n".join(lines).strip()
